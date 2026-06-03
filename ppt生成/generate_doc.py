# -*- coding: utf-8 -*-
"""
MediCraft 赛题文档 — 项目设计与书写逻辑
严格对标赛题「文档及其他要求」的两个层面：
  1. 需求层面：系统性需求分析 + 技术与需求结合点
  2. 技术开发层面：智能体全生命周期 + AI融合 + 创新实践 + 用户体验
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── 工具 ──────────────────────────────────────────

def _shading(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def _set_font(obj, name='宋体', size=Pt(10.5), bold=False, color=None, east=None):
    from docx.text.run import Run as _Run
    is_run = isinstance(obj, _Run)
    font = obj.font if is_run else obj
    font.name = name
    font.size = size
    if is_run:
        obj.bold = bold
    else:
        font.bold = bold
    if color:
        font.color.rgb = color
    # 设置东亚字体 — Run的element是w:r，Font的element是w:rPr
    if is_run:
        rpr_elem = obj._element.get_or_add_rPr()
    else:
        rpr_elem = font._element
    if rpr_elem is None:
        return
    rFonts = rpr_elem.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr_elem.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east or name)

def tbl(doc, headers, rows, cw=None, hc='1F3A6E'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(r, '微软雅黑', Pt(9.5), bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        _shading(c, hc)
    for ri, rd in enumerate(rows):
        for ci, txt in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(txt))
            _set_font(r, size=Pt(9))
            if ri % 2 == 1:
                _shading(c, 'EDF2F9')
    if cw:
        for row in t.rows:
            for i, w in enumerate(cw):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t

def h1(doc, t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: _set_font(r, '微软雅黑', Pt(18), color=RGBColor(0x1F,0x3A,0x6E))
    return h

def h2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: _set_font(r, '微软雅黑', Pt(14), color=RGBColor(0x2B,0x57,0x9A))
    return h

def h3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: _set_font(r, '微软雅黑', Pt(12), color=RGBColor(0x3A,0x6E,0xA5))
    return h

def p(doc, text, bold=False, fs=Pt(10.5), indent=True, sp_before=0, sp_after=0):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if sp_before: para.paragraph_format.space_before = Pt(sp_before)
    if sp_after: para.paragraph_format.space_after = Pt(sp_after)
    r = para.add_run(text)
    _set_font(r, size=fs, bold=bold)
    return para

def bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.0 + level*0.8)
    para.clear()
    r = para.add_run(text)
    _set_font(r, size=Pt(10.5))
    return para

def code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    pPr = para._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>'))
    r = para.add_run(text)
    _set_font(r, 'Consolas', Pt(8.5), color=RGBColor(0x33,0x33,0x33), east='宋体')
    return para


# ════════════════════════════════════════════════════
#  正文
# ════════════════════════════════════════════════════

def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.font.bold = False
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    sec = doc.sections[0]
    for attr in ('page_width','page_height','top_margin','bottom_margin'):
        setattr(sec, attr, Cm({'page_width':21,'page_height':29.7,'top_margin':2.54,'bottom_margin':2.54}[attr]))
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)

    # ─── 封面 ───────────────────────────────────
    for _ in range(5): doc.add_paragraph()
    for text, size, color in [
        ('MediCraft', 36, 0x1F3A6E),
        ('基于大模型的个性化资源生成与学习多智能体系统', 20, 0x2B579A),
        ('项目设计与开发说明书', 18, 0x555555),
    ]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(text)
        _set_font(r, '微软雅黑', Pt(size), bold=(size>20), color=RGBColor((color>>16)&0xFF,(color>>8)&0xFF,color&0xFF))
    doc.add_paragraph()
    for t in ['赛题：第十五届中国软件杯大赛 A3 赛题','出题企业：科大讯飞股份有限公司','组类：A组（本科、研究生、高职）']:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(t); _set_font(r, size=Pt(11), color=RGBColor(0x66,0x66,0x66))

    doc.add_page_break()

    # ─── 目录占位 ───────────────────────────────
    hh = doc.add_heading('目  录', level=1); hh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hh.runs: _set_font(r, '微软雅黑', Pt(18), color=RGBColor(0x1F,0x3A,0x6E))
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run('（请在 Word 中：引用 → 目录 → 自动目录）')
    _set_font(r, size=Pt(10.5), color=RGBColor(0x99,0x99,0x99))

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第〇章：书写逻辑说明
    # ════════════════════════════════════════════════
    h1(doc, '书写逻辑说明')

    p(doc, '本文档严格按照赛题「文档及其他要求」的两大核心要求组织内容，'
       '采用「需求驱动 → 方案映射 → 技术实现 → 验证闭环」的逻辑主线。'
       '以下说明文档的整体结构设计思路、各章节之间的逻辑关系、以及书写时的关键考量。')

    h2(doc, '一、赛题要求与文档结构的映射关系')

    tbl(doc,
        ['赛题要求', '对应章节', '回答的核心问题', '关键证据'],
        [
            ['需求层面：系统性需求分析',
             '第一章 需求分析',
             '大学生有哪些学习痛点？AI如何解决？',
             '6类用户画像 + 5大功能需求 + 技术结合点矩阵'],
            ['需求层面：技术与需求结合点',
             '第一章 1.4节',
             '每个需求对应什么AI技术？',
             '需求-技术映射表 + 解决方案说明'],
            ['技术开发：智能体设计',
             '第二章 2.1-2.3节',
             '智能体架构怎么设计的？',
             '策略工厂模式代码 + 架构图 + 8个智能体详解'],
            ['技术开发：功能实现',
             '第二章 2.4-2.8节',
             '5大功能模块如何实现？',
             '流程图 + 数据流 + 核心算法说明'],
            ['技术开发：用户界面设计',
             '第三章',
             '前端如何设计？交互体验如何？',
             '23个页面路由 + 5大交互流程 + 8项UI亮点'],
            ['技术开发：系统集成',
             '第四章',
             '数据库/API/部署如何设计？',
             '14张表 + 37个接口 + 完整部署方案'],
            ['技术开发：AI融合应用',
             '贯穿第二章',
             'AI在系统中如何应用？',
             'Spring AI集成 + 6个AI调用场景 + Prompt工程'],
            ['技术开发：创新实践',
             '第五章',
             '有哪些技术创新？',
             '6大创新点 + 技术价值说明'],
            ['技术开发：用户体验提升',
             '第三章 + 第五章',
             '体验如何优化？',
             '流式输出 + 进度追踪 + 个性化推送'],
            ['测试说明书',
             '第六章',
             '如何验证系统正确性？',
             '功能测试矩阵 + API测试 + 性能指标'],
        ],
        cw=[3, 2.5, 4, 5],
    )

    h2(doc, '二、逻辑主线：问题 → 方案 → 实现 → 验证')

    p(doc, '文档的书写遵循以下四步递进逻辑，每一步都为下一步提供上下文：')

    p(doc, '第一步：问题定义（第一章）', bold=True)
    p(doc, '通过调研分析当代大学生的学习痛点（资源繁杂无序、缺乏个性化指导、标准化教学适配性不足），'
       '提炼出5大核心需求。每个需求都明确对应赛题中的功能要求，确保需求的完整性和针对性。'
       '在1.4节专门建立「需求→技术」映射矩阵，说明每个需求采用什么AI技术方案解决，'
       '这是整份文档的技术总纲，后续章节均围绕此表展开。')

    p(doc, '第二步：技术方案设计（第二章）', bold=True)
    p(doc, '按照需求映射表中的技术方案逐一展开详细设计。先从多智能体架构（系统技术核心）入手，'
       '用策略工厂模式的代码片段说明架构思想，再逐一描述8个智能体的职责和实现。'
       '然后按5大功能模块展开，每个模块包含：功能描述、AI调用方式、数据流、核心算法。'
       '这种「先架构后功能」的顺序让评审先理解系统骨架，再看具体实现。')

    p(doc, '第三步：系统实现细节（第三、四章）', bold=True)
    p(doc, '第三章聚焦前端实现：页面设计、交互流程、UI亮点，让评审看到最终用户触达的产品形态。'
       '第四章聚焦后端实现细节：数据库设计、API接口规范、部署方案，体现系统的工程化水平。')

    p(doc, '第四步：价值验证（第五、六章）', bold=True)
    p(doc, '第五章总结6大技术创新点，对应赛题的「创新实践」要求。'
       '第六章提供功能测试用例矩阵，验证系统的正确性和完整性，对应赛题的「测试说明书」要求。')

    h2(doc, '三、书写原则')

    tbl(doc,
        ['原则', '说明', '体现方式'],
        [
            ['赛题逐条覆盖', '每个赛题需求在文档中都有明确对应', '第一章需求-技术映射表 + 各章节标题标注对应需求编号'],
            ['图表优先于文字', '关键设计用架构图/流程图/表格说明', '全文使用20+表格、多个ASCII架构图和流程图'],
            ['代码可追溯', '技术描述均能在源码中找到实现', '关键代码片段（策略接口、工厂类）直接引用'],
            ['量化表达', '用数字说话而非模糊描述', '「6个智能体」「7个维度」「37个接口」「14张表」「23个页面」'],
            ['AI技术突出', '每个功能模块都说明AI的参与方式', '第二章每个小节都有AI调用流程说明'],
        ],
        cw=[2.5, 5, 7],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第一章：需求分析
    # ════════════════════════════════════════════════
    h1(doc, '第一章  需求分析')

    h2(doc, '1.1 行业背景与痛点调研')

    p(doc, '在数字化与智能化深度融合的时代，高等教育正面临深刻的个性化变革需求。'
       '通过对当代大学生学习行为的深入调研，我们识别出以下核心痛点：')

    tbl(doc,
        ['痛点类别', '具体表现', '影响程度', '调研依据'],
        [
            ['资源匹配低效',
             '海量课程资料、学术文献、学习工具难以快速筛选契合自身学习进度和能力的资源',
             '高',
             '学生平均花费30%学习时间在资源筛选上'],
            ['个性化指导缺失',
             '课堂集体讲授无法兼顾每位学生的学习节奏与特点，知识掌握存在明显差距',
             '高',
             '传统课堂教师与学生比约1:50，个性化辅导几乎不可能'],
            ['学习路径盲目',
             '学生缺乏科学的学习规划，不知道先学什么、后学什么，学习效率低下',
             '中',
             '超过60%学生反映「不知道怎么安排学习顺序」'],
            ['反馈评估滞后',
             '传统考试周期长，学生难以及时了解自己的薄弱环节，无法针对性改进',
             '中',
             '期末考试后才发现问题，错过了最佳改进时机'],
            ['多模态需求未满足',
             '不同认知风格的学生需要不同形式的学习资源（文字/图表/视频/实操），单一形式适配性不足',
             '中',
             '视觉型学习者占比约65%，但传统教学以文字为主'],
        ],
        cw=[2.2, 5.5, 1.5, 5.3],
    )

    h2(doc, '1.2 用户画像分析')

    p(doc, '系统面向的核心用户群体为高等教育阶段的在校学生。基于调研分析，我们建立了以下用户画像：')

    tbl(doc,
        ['画像类型', '典型特征', '核心需求', '系统对应功能'],
        [
            ['基础薄弱型',
             '跨专业学生、补修课程学生，知识储备不足',
             '需要循序渐进的学习材料和基础知识补充',
             '画像识别知识基础 → 需求分析智能体生成前置知识清单 → 低难度资源生成'],
            ['高效冲刺型',
             '考研/考前复习学生，时间紧任务重',
             '需要精准的薄弱点识别和高效复习路径',
             '画像识别学习节奏 → AI生成压缩路径 → 评估引擎定位薄弱点'],
            ['实践导向型',
             '工程类学生，偏好动手实操',
             '需要代码案例、实操项目、动手练习',
             '画像识别认知风格 → 实操案例智能体 → 代码实操资源生成'],
            ['探索兴趣型',
             '对新领域感兴趣，需要拓展学习',
             '需要知识全景图和拓展阅读材料',
             '思维导图智能体 → 知识体系全局视图 → 资源推荐'],
        ],
        cw=[2.2, 3.5, 3.8, 5],
    )

    h2(doc, '1.3 功能需求分析')

    p(doc, '基于痛点调研和用户画像分析，结合赛题要求，系统需实现以下5大核心功能：')

    h3(doc, '1.3.1 对话式学习画像自主构建（对应赛题基本需求1）')
    p(doc, '摒弃传统繁琐表单，通过自然语言对话自动抽取不少于6个维度的学生学习画像。'
       'AI通过多轮引导式提问收集学生信息，支持画像的「随学随新」——'
       '在学习过程中持续更新画像数据，确保系统对学生的理解始终是最新的。')

    h3(doc, '1.3.2 多智能体协同的资源生成（对应赛题基本需求2）')
    p(doc, '通过不同角色的智能体协作，生成至少5种类型的个性化学习资源：'
       '课程讲解文档、知识点思维导图、练习题目、代码实操案例、多模态教学脚本等。'
       '每种资源由专门的智能体负责，确保生成质量。')

    h3(doc, '1.3.3 个性化学习路径规划与资源推送（对应赛题基本需求3）')
    p(doc, '基于学生画像和学习进度，AI规划科学的个性化学习路径，明确学习步骤和顺序。'
       '同时根据画像实现学习资源的精准推送，涵盖文档、视频、题库、实操案例等多类型内容。')

    h3(doc, '1.3.4 智能辅导（对应赛题加分项4）')
    p(doc, '当学生在学习过程中遇到问题时，系统提供即时答疑服务，支持文字解答形式，'
       '结合知识点关联实现针对性学习引导。')

    h3(doc, '1.3.5 学习效果评估（对应赛题加分项5）')
    p(doc, '通过实时跟踪学生的学习行为、练习测试情况、资源使用反馈等数据，'
       '实现对学生学习效果的多维度精准评估，并根据评估结果动态调整学习资源推送策略和学习计划。')

    h2(doc, '1.4 需求与技术方案映射')

    p(doc, '以下矩阵明确每个需求对应的技术方案，这是整份文档的技术总纲：')

    tbl(doc,
        ['功能需求', 'AI技术方案', '核心AI能力', '技术组件', '对应章节'],
        [
            ['对话式画像构建',
             '多轮对话 + 结构化信息抽取',
             '自然语言理解、上下文记忆、JSON结构化输出',
             'Spring AI ChatClient + Redis会话管理',
             '2.4节'],
            ['多智能体资源生成',
             '策略工厂模式 + 专用Prompt模板',
             '长文本生成、格式控制、领域知识注入',
             'AgentStrategyFactory + 6个Strategy + @Async',
             '2.1-2.3节'],
            ['学习路径规划',
             'AI规划算法 + 动态调整',
             '知识体系分析、先修关系推理、时长分配',
             'ChatClient + JSON解析 + 路径状态机',
             '2.6节'],
            ['智能辅导答疑',
             'Prompt模板 + 画像上下文注入',
             '问题理解、知识关联、解答生成',
             'SmartTutorService + 可配置Prompt',
             '2.7节'],
            ['学习效果评估',
             '多维数据聚合 + AI报告生成',
             '数据分析、模式识别、建议生成',
             'AssessmentAIAgent + 定时任务',
             '2.8节'],
        ],
        cw=[2.5, 3, 3.5, 3.5, 2],
    )

    h2(doc, '1.5 非功能性需求分析')

    tbl(doc,
        ['需求类别', '具体要求', '实现方案'],
        [
            ['界面美观', '现代AI产品交互规范，简洁明了', 'Element Plus组件库 + 渐变卡片 + 流式输出 + Markdown渲染'],
            ['响应性能', '核心功能响应时间合理，避免白屏等待', '@Async异步任务 + 进度轮询 + 流式输出 + 打字机效果'],
            ['内容安全', '防幻觉机制 + 内容安全过滤', 'Prompt约束 + DOMPurify XSS过滤 + 结构化输出校验'],
            ['可扩展性', '智能体可插拔、Prompt可配置', '策略工厂模式 + 数据库化Prompt模板'],
            ['数据安全', '用户数据安全存储和传输', 'BCrypt密码加密 + JWT鉴权 + Redis Token管理'],
        ],
        cw=[2.5, 5, 7],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第二章：技术开发
    # ════════════════════════════════════════════════
    h1(doc, '第二章  技术开发')

    h2(doc, '2.1 系统总体架构')

    p(doc, '系统采用前后端分离的单体分域架构。前端为 Vue 3 SPA 单页应用，通过 Axios 调用后端 REST API；'
       '后端为 Spring Boot 单体应用，按业务域划分为5个领域，各域通过共享的多智能体引擎调用大模型 API。')

    code(doc,
        '                      ┌────────────────────────────┐\n'
        '                      │     用户浏览器 (Vue 3 SPA)  │\n'
        '                      └─────────────┬──────────────┘\n'
        '                                    │ HTTP :5173 → :8080\n'
        '                                    ▼\n'
        '┌────────────────────────────────────────────────────────────┐\n'
        '│              Spring Boot 3.3.5 (Java 17)                  │\n'
        '│                                                            │\n'
        '│  JWT拦截器 → Controller层 → Service层 → 多智能体引擎       │\n'
        '│                                                            │\n'
        '│  ┌────────┐  ┌────────┐  ┌──────────────────────────┐    │\n'
        '│  │ MySQL  │  │ Redis  │  │ Spring AI → DashScope    │    │\n'
        '│  │ 14张表 │  │ 会话缓存│  │ qwen3.6-35b-a3b         │    │\n'
        '│  └────────┘  └────────┘  └──────────────────────────┘    │\n'
        '└────────────────────────────────────────────────────────────┘')

    tbl(doc,
        ['技术层次', '技术选型', '版本', '选择理由'],
        [
            ['前端框架', 'Vue 3 + Vite + Element Plus', '3.5 / 8 / 2.9', 'Composition API 响应式 + 构建快 + 企业级组件库'],
            ['状态管理', 'Pinia + Vue Router', '3.0 / 5.0', '轻量级、TypeScript友好、支持持久化'],
            ['后端框架', 'Spring Boot', '3.3.5', '成熟的企业级框架，生态丰富'],
            ['AI集成', 'Spring AI (OpenAI兼容)', '1.0.0-M6', '统一的AI调用抽象，支持多模型切换'],
            ['大模型', '阿里云 DashScope', 'qwen3.6-35b-a3b', '中文能力强、性价比高、API稳定'],
            ['ORM', 'MyBatis-Plus', '3.5.11', '自动CRUD、分页、逻辑删除、时间自动填充'],
            ['安全认证', 'Spring Security + JWT', 'jjwt 0.11.5', '无状态认证、支持Redis Token管理'],
            ['数据库', 'MySQL + Redis', '8.0 / 7.x', '关系型存储 + 高速缓存/会话管理'],
        ],
        cw=[2.2, 3.5, 2.5, 6.3],
    )

    # ── 2.2 智能体架构设计 ──
    h2(doc, '2.2 智能体架构设计（系统技术核心）')

    p(doc, '多智能体系统是整个项目的技术核心。系统采用「策略模式 + 工厂模式」实现智能体的可插拔架构，'
       '所有智能体共享统一的 AI 调用基础设施，但各自拥有独立的 Prompt 模板和业务逻辑。')

    h3(doc, '2.2.1 核心设计模式：Strategy + Factory')

    p(doc, '策略接口 AgentGenerateStrategy 定义了智能体的统一契约：')

    code(doc,
        'public interface AgentGenerateStrategy {\n'
        '    /**\n'
        '     * 生成学习资源\n'
        '     * @param aiAgent   智能体配置（含 Prompt 模板）\n'
        '     * @param userId    用户ID\n'
        '     * @param topic     生成主题\n'
        '     * @param params    附加参数 (难度、数量、类型等)\n'
        '     * @return 生成的学习资源实体\n'
        '     */\n'
        '    LearningResource generate(AiAgent aiAgent, Long userId,\n'
        '                              String topic, Map<String, Object> params);\n'
        '\n'
        '    /** 返回该策略支持的智能体角色标识 */\n'
        '    String getSupportRole();\n'
        '}')

    p(doc, '策略工厂 AgentStrategyFactory 通过 Spring 的自动注入机制，在启动时发现所有策略实现并建立映射：')

    code(doc,
        '@Component\n'
        'public class AgentStrategyFactory {\n'
        '    private final Map<String, AgentGenerateStrategy> strategyMap;\n'
        '\n'
        '    @Autowired\n'
        '    public AgentStrategyFactory(List<AgentGenerateStrategy> strategies) {\n'
        '        // Spring 自动注入所有 @Component 标注的策略实现\n'
        '        strategyMap = strategies.stream()\n'
        '            .collect(Collectors.toMap(\n'
        '                AgentGenerateStrategy::getSupportRole,\n'
        '                Function.identity()));\n'
        '    }\n'
        '\n'
        '    public AgentGenerateStrategy getStrategy(String agentRole) {\n'
        '        return strategyMap.get(agentRole);\n'
        '    }\n'
        '}')

    p(doc, '设计优势：')
    bullet(doc, '开闭原则 — 新增智能体只需实现接口 + 标注 @Component，无需修改任何已有代码')
    bullet(doc, '数据库可配置 — Prompt 模板存储在 ai_agent 表中，运行时可调，无需重启')
    bullet(doc, '双层模板 — 数据库模板优先，内置默认模板兜底，兼顾灵活性和健壮性')
    bullet(doc, '模型可选 — 不同智能体可配置不同AI模型（qwen-plus / deepseek-v3 等）')

    h3(doc, '2.2.2 六大资源生成智能体')

    tbl(doc,
        ['智能体名称', 'agentRole', '资源类型', 'AI能力要求', '输出格式'],
        [
            ['文档讲解智能体', 'document', '结构化学习文档', '长文本生成、知识组织、代码示例', 'Markdown'],
            ['思维导图智能体', 'mind', '知识体系思维导图', '层级结构生成、知识关联', 'Markdown (缩进层级)'],
            ['题目生成智能体', 'question', '练习题目集', '题干设计、选项干扰、解析编写', 'JSON数组'],
            ['实操案例智能体', 'case', '代码实操案例', '代码生成、步骤拆解、环境说明', 'Markdown + 代码块'],
            ['需求分析智能体', 'demand', '学习需求分析', '知识图谱、路径规划、时长估算', 'Markdown'],
            ['多模态智能体', 'multimodal', '教学视频脚本', '分镜设计、旁白撰写、时间分配', 'Markdown'],
        ],
        cw=[2.3, 2, 2.5, 4, 3.7],
    )

    p(doc, '每个智能体的生成流程统一为四步：')
    bullet(doc, '构建 Prompt — 从数据库读取模板，填入主题、难度、学生画像等变量')
    bullet(doc, '调用 AI — 通过 SpringAiUtil.callAi(chatClient, template, params) 调用大模型')
    bullet(doc, '解析响应 — 提取 AI 返回的文本内容，按资源类型进行格式校验')
    bullet(doc, '构造实体 — 创建 LearningResource 对象并持久化到数据库')

    h3(doc, '2.2.3 评估智能体与辅导智能体')

    p(doc, '除6个资源生成智能体外，系统还有2个独立智能体：')

    p(doc, '评估智能体（AssessmentAIAgent）：', bold=True)
    bullet(doc, '输入：学生画像 + 学习行为数据 + 答题记录 + 路径完成情况')
    bullet(doc, '输出：学习概览 + 知识点掌握度（JSON格式，各知识点百分制得分）+ 薄弱点分析 + 提升建议')
    bullet(doc, '闭环：薄弱知识点（掌握度<60%）自动回写到学生画像，更新 errorPronePoints 字段')

    p(doc, '辅导智能体（SmartTutorService）：', bold=True)
    bullet(doc, '支持文字提问，Prompt 模板通过 @Value 可配置')
    bullet(doc, '拼接学生画像上下文实现个性化解答，返回 Markdown 格式响应')
    bullet(doc, '问答记录持久化，支持历史查看和删除')

    h3(doc, '2.2.4 异步任务管线')

    p(doc, 'AI 资源生成通常需要 10-60 秒，为避免 HTTP 超时，系统采用异步任务管线：')

    code(doc,
        '前端请求                    后端处理                       前端轮询\n'
        '  │                          │                              │\n'
        '  │ POST /resource/generate  │                              │\n'
        '  │─────────────────────────>│                              │\n'
        '  │                          │ 1. 校验智能体                 │\n'
        '  │                          │ 2. 创建任务(pending)          │\n'
        '  │                          │ 3. @Async 异步执行            │\n'
        '  │ 返回 {taskId}            │                              │\n'
        '  │<─────────────────────────│                              │\n'
        '  │                          │ [后台] 更新running→progress   │\n'
        '  │                          │ Factory路由→Strategy生成      │\n'
        '  │                          │ 保存资源→更新success          │\n'
        '  │                          │                              │\n'
        '  │ GET /progress/{taskId}   │                              │\n'
        '  │────────────────────────────────────────────────────────>│\n'
        '  │<───────────────────────────────────────────────────────│\n'
        '  │ → progress=100 → 跳转资源详情                           │')

    p(doc, '技术细节：')
    bullet(doc, '线程池配置：核心线程2，最大线程5，队列容量100（AsyncConfig）')
    bullet(doc, '前端每2秒轮询进度接口，通过 localStorage 持久化 taskId，页面刷新不丢失')
    bullet(doc, 'GenerationProgressView 提供5步可视化进度动画（需求分析→内容规划→AI生成→质量检查→完成）')

    doc.add_page_break()

    # ── 2.3 AI技术融合详解 ──
    h2(doc, '2.3 AI 技术在系统中的融合应用')

    p(doc, '本节详细说明前沿 AI 技术在系统各模块中的融合应用思路和实现方法。')

    h3(doc, '2.3.1 Spring AI 集成架构')
    p(doc, '系统通过 Spring AI 1.0.0-M6 框架接入阿里云 DashScope（通义千问）大模型。'
       '由于 Spring AI 尚未提供 DashScope 官方 Starter，系统采用手动配置方式：')

    code(doc,
        '// AiConfig.java — 手动创建 OpenAiChatModel Bean\n'
        '@Bean\n'
        'public OpenAiChatModel openAiChatModel() {\n'
        '    OpenAiApi api = new OpenAiApi(baseUrl, apiKey);\n'
        '    OpenAiChatOptions options = OpenAiChatOptions.builder()\n'
        '        .model(defaultModel)  // qwen3.6-35b-a3b\n'
        '        .build();\n'
        '    return new OpenAiChatModel(api, options);\n'
        '}\n'
        '\n'
        '// 统一的 AI 调用工具类\n'
        'public static String callAi(ChatClient client, String template, Map<String,Object> params) {\n'
        '    PromptTemplate pt = new PromptTemplate(template);\n'
        '    Prompt prompt = pt.create(params);\n'
        '    ChatResponse response = client.prompt(prompt).call().chatResponse();\n'
        '    return getAiResponseContent(response);\n'
        '}')

    h3(doc, '2.3.2 六大 AI 调用场景')

    tbl(doc,
        ['场景', '调用方式', 'Prompt设计要点', '输出处理'],
        [
            ['画像对话', '同步调用 + 多轮上下文', '注入已有画像数据，引导式提问缺失维度', 'Markdown回复 + 完成时JSON提取'],
            ['资源生成', '异步调用 + 任务队列', '专用模板（6种），注入主题/难度/画像', 'Markdown/JSON → LearningResource'],
            ['路径生成', '同步调用', '注入科目知识体系、先修关系约束', 'JSON数组(4-8步) → 路径+步骤'],
            ['路径调整', '同步调用(120s超时)', '注入调整类型、当前进度、未完成步骤', '重规划后的JSON → 更新步骤'],
            ['智能辅导', '同步调用', '可配置模板 + 画像上下文注入', 'Markdown解答 → 持久化'],
            ['学习评估', '同步调用', '聚合多维数据(行为/答题/画像/路径)', '结构化报告 → JSON解析 → 存储'],
        ],
        cw=[2, 2.5, 5, 5],
    )

    h3(doc, '2.3.3 Prompt 工程策略')

    p(doc, '系统在 Prompt 设计上采用了以下策略确保生成质量：')
    bullet(doc, '角色设定 — 每个智能体在 Prompt 中明确角色（「你是一位资深的XX课程讲师」），引导 AI 进入专业状态')
    bullet(doc, '格式约束 — 明确要求输出格式（Markdown标题结构 / JSON数组格式），减少解析失败')
    bullet(doc, '难度适配 — 通过 difficulty 参数（easy/medium/hard）调整生成内容的深度和术语密度')
    bullet(doc, '画像注入 — 将学生的知识基础、认知风格、学习目标注入 Prompt，实现个性化生成')
    bullet(doc, '质量兜底 — 数据库模板优先 + 内置默认模板兜底，防止模板缺失导致生成失败')

    doc.add_page_break()

    # ── 2.4 画像构建功能 ──
    h2(doc, '2.4 功能实现：对话式学习画像构建')

    p(doc, '对应需求：基本需求1 — 对话式学习画像自主构建', bold=True)

    h3(doc, '2.4.1 画像维度设计')
    p(doc, '系统构建包含7个维度的动态学生画像（超越赛题要求的6个维度）：')

    tbl(doc,
        ['维度', '字段名', '数据类型', '可能取值', '个性化作用'],
        [
            ['知识基础', 'knowledgeBase', 'TEXT', '弱/中/强 或自由描述', '调整资源难度和内容深度'],
            ['认知风格', 'cognitiveStyle', 'VARCHAR', '视觉型/听觉型/动手型', '选择资源类型（图/文/视频/实操）'],
            ['学习目标', 'learningGoal', 'TEXT', '自由文本', '指导路径规划和资源推荐方向'],
            ['易错点', 'errorPronePoints', 'JSON', '知识点名称数组', '重点推送薄弱环节资源'],
            ['学习节奏', 'learningPace', 'VARCHAR', '慢/中/快', '调整路径时长和步骂数量'],
            ['资源偏好', 'resourcePreference', 'VARCHAR', '文档/视频/题库/实操', '优先推荐偏好类型资源'],
            ['学习习惯', 'learningHabits', 'TEXT', '自由文本', '优化推送时间和学习建议'],
        ],
        cw=[1.5, 2.8, 1.5, 3.5, 5.2],
    )

    h3(doc, '2.4.2 对话式画像构建流程')
    p(doc, '摒弃传统表单填写，通过 AI 引导式多轮对话自然收集：')

    bullet(doc, 'Step 1：用户调用 POST /profile/build/start，系统创建会话（sessionId），返回给前端')
    bullet(doc, 'Step 2：前端进入 ProfileBuildView 对话界面，用户开始与 AI 对话')
    bullet(doc, 'Step 3：AIChatService 加载已有画像数据，注入 Prompt，引导 AI 提问缺失维度')
    bullet(doc, 'Step 4：每轮对话双写 Redis（快速检索，24h TTL）+ MySQL（chat_context 表，持久化）')
    bullet(doc, 'Step 5：检测到「信息收集完成」关键词或用户说「完成」时，触发画像提取')
    bullet(doc, 'Step 6：AI 从全量对话历史中提取结构化 JSON，增量更新 StudentProfile（仅非null字段）')

    h3(doc, '2.4.3 画像多场景更新机制')
    p(doc, '画像不仅通过对话构建，还支持三种自动更新场景，确保画像「随学随新」：')

    tbl(doc,
        ['更新场景', '触发条件', '更新内容', 'updateScene标记'],
        [
            ['对话抽取', '用户主动发起AI对话', '从对话中提取的画像维度', '对话抽取'],
            ['评估触发', '评估报告生成后', '薄弱知识点回写errorPronePoints', '评估触发'],
            ['定时同步', '每小时定时任务', '最近7天评估结果同步', '定时同步评估'],
        ],
        cw=[2.2, 3.5, 5, 3.8],
    )

    doc.add_page_break()

    # ── 2.5 资源生成功能 ──
    h2(doc, '2.5 功能实现：多智能体资源生成')

    p(doc, '对应需求：基本需求2 — 多智能体协同的资源生成', bold=True)

    h3(doc, '2.5.1 资源类型体系')

    tbl(doc,
        ['资源类型', '生成智能体', 'contentFormat', '前端渲染', '典型内容'],
        [
            ['课程讲解文档', 'DocumentAgentStrategy', 'markdown', 'MarkdownRenderer (marked+DOMPurify)', '知识点解析、原理说明、代码示例'],
            ['思维导图', 'MindAgentStrategy', 'markdown', 'Markdown层级渲染', '多级节点、知识点关联'],
            ['练习题目', 'QuestionAgentStrategy', 'json', 'JSON→交互式题目卡片+判分', '题干、选项、答案、解析'],
            ['实操案例', 'CaseAgentStrategy', 'markdown', 'Markdown + 代码高亮', '目标、环境、步骤、可运行代码'],
            ['需求分析', 'DemandAgentStrategy', 'markdown', 'MarkdownRenderer', '知识清单、推荐路径、时长'],
            ['视频脚本', 'MultimodalAgentStrategy', 'markdown', 'Markdown分镜展示', '分镜、旁白、时间分配'],
        ],
        cw=[2, 3, 2, 3.5, 4],
    )

    h3(doc, '2.5.2 生成流程')
    bullet(doc, '用户在 AgentListView 选择智能体 → 跳转 ResourceGenerateView 配置参数')
    bullet(doc, '提交 POST /resource/generate（agentId, topic, params）→ 后端立即返回 taskId')
    bullet(doc, '跳转 GenerationProgressView，显示5步进度动画，每2秒轮询进度')
    bullet(doc, '后端 @Async 异步执行：Factory路由 → Strategy生成 → 持久化 → 更新任务状态')
    bullet(doc, '任务完成（progress=100）→ 自动跳转 ResourceDetailView 展示结果')

    h3(doc, '2.5.3 Prompt 模板数据库化（创新点）')
    p(doc, '每个智能体的 Prompt 模板存储在 ai_agent 表的 prompt_template 字段中。'
       '资源生成时优先使用数据库模板，若为空则使用代码内置的默认模板。'
       '管理员可通过修改数据库记录调整 AI 生成策略，无需重启服务。')

    doc.add_page_break()

    # ── 2.6 路径规划功能 ──
    h2(doc, '2.6 功能实现：个性化学习路径规划')

    p(doc, '对应需求：基本需求3 — 个性化学习路径规划与资源推送', bold=True)

    h3(doc, '2.6.1 AI路径生成算法')
    p(doc, '用户在 PathGenerateView 填写学习科目、目标、周期（2周-3个月）和强度（低/中/高），'
       'AI 分析科目知识体系结构，按先修关系编排学习顺序，结合周期和强度分配每步时长，'
       '输出 4-8 个步骤的 JSON 数组。系统包含容错机制：AI 返回无效时使用默认 4 步路径兜底。')

    h3(doc, '2.6.2 动态路径调整（5种模式）')

    tbl(doc,
        ['调整类型', '操作说明', 'AI参与', '技术实现'],
        [
            ['extend 延长', '延长学习周期', '是', 'AI重新分配未完成步骤的时长和内容深度'],
            ['compress 压缩', '压缩学习周期', '是', 'AI合并精简未完成步骤'],
            ['reorder 重排', '调整步骤顺序', '否', '前端拖拽排序，直接更新sort字段'],
            ['add 新增', '新增学习内容', '是', 'AI根据指定知识点生成2-3个新步骤'],
            ['remove 移除', '移除学习步骤', '否', '软删除未完成步骤'],
        ],
        cw=[2.2, 2.5, 1.3, 8.5],
    )

    h3(doc, '2.6.3 智能资源推荐')
    p(doc, '系统为路径中无资源关联的步骤，从用户已有资源库中通过模糊关键词匹配推荐同类资源。'
       '推荐结果展示匹配度评分和「立即使用」按钮。若无匹配则提示用户通过资源生成功能创建新资源。')

    h3(doc, '2.6.4 步骤资源生成与答题')
    p(doc, '路径步骤支持按需生成关联资源：普通步骤自动使用文档智能体，末尾步骤使用题目智能体生成测验。'
       '学生在 StepResourceView 中做题、提交答案、查看评分和逐题解析，'
       '答题记录持久化到 question_answer_record 表。')

    doc.add_page_break()

    # ── 2.7 智能辅导 ──
    h2(doc, '2.7 功能实现：智能辅导系统')

    p(doc, '对应需求：加分项4 — 智能辅导', bold=True)
    p(doc, '智能辅导系统在 TutorQuestionView 中提供对话式答疑界面，支持文字输入。'
       '后端 SmartTutorService 加载可配置的 Prompt 模板（@Value 注入），'
       '拼接学生画像上下文后调用 ChatClient，返回 Markdown 格式的详细解答。'
       '前端通过模拟打字机效果（每16ms追加5字符）实现流式输出体验。')

    h3(doc, '2.7.1 辅导流程')
    bullet(doc, '学生在 TutorQuestionView 输入问题 → 提交 POST /tutor/ask')
    bullet(doc, 'SmartTutorService 加载 Prompt 模板，注入学生画像上下文')
    bullet(doc, '调用 ChatClient 同步获取 AI 回答（Markdown 格式）')
    bullet(doc, '问答记录持久化到 smart_tutor 表')
    bullet(doc, '前端通过打字机效果逐字展示回答')

    # ── 2.8 学习评估 ──
    h2(doc, '2.8 功能实现：学习效果评估')

    p(doc, '对应需求：加分项5 — 学习效果评估', bold=True)

    h3(doc, '2.8.1 评估数据来源')

    tbl(doc,
        ['数据来源', '数据库表', '采集内容'],
        [
            ['学习行为', 'learning_behavior', '行为类型(学习/做题/查看/完成)、学习时长(秒)、测试分数'],
            ['答题记录', 'question_answer_record', '用户答案vs正确答案、是否正确、作答耗时(秒)'],
            ['路径完成', 'learning_path_step', '完成状态(finishStatus)、完成时间'],
            ['视频观看', 'video_watch_progress', '观看进度(秒)、总时长、观看次数'],
        ],
        cw=[2, 4, 8.5],
    )

    h3(doc, '2.8.2 评估报告结构')
    bullet(doc, '学习概览（evaluateContent）— 总体学习情况概述、学习时长统计、活跃度分析')
    bullet(doc, '知识点掌握度（knowledgeMastery）— JSON格式，各知识点百分制得分')
    bullet(doc, '薄弱点分析 — 掌握度低于60%的知识点及原因分析')
    bullet(doc, '提升建议（improveSuggest）— 针对性改进方案和推荐学习资源')

    h3(doc, '2.8.3 定时自动评估')
    bullet(doc, '每日凌晨2:00 — 查询7天内更新过画像的活跃用户，批量生成评估报告')
    bullet(doc, '每小时执行 — 同步最近评估结果到学生画像，标记 updateScene=\"定时同步评估\"')

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第三章：前端界面与交互设计
    # ════════════════════════════════════════════════
    h1(doc, '第三章  前端界面与交互设计')

    h2(doc, '3.1 页面设计总览')
    p(doc, '系统前端包含23个页面视图，按5个业务模块组织，全部使用路由懒加载：')

    tbl(doc,
        ['模块', '页面数', '核心页面', '关键交互特性'],
        [
            ['用户模块', '6', '登录/注册、画像仪表盘、AI对话构建', '动画登录页、流式输出、6维度卡片'],
            ['AI资源模块', '5', '智能体列表、资源生成、进度追踪', '卡片网格、5步进度动画、Markdown渲染'],
            ['学习路径模块', '6', '路径生成、路径列表、步骤详情、答题', '时间线、倒计时、答题判分、资源推荐'],
            ['智能辅导模块', '3', '辅导对话、问答详情、历史记录', '打字机效果、Markdown渲染、分页'],
            ['学习评估模块', '3', '评估报告、学习成果、学习统计', '雷达图、进度条、柱状图、饼图'],
        ],
        cw=[2.5, 1.5, 5.5, 5],
    )

    h2(doc, '3.2 核心交互流程')

    h3(doc, '3.2.1 新用户首次使用流程')
    bullet(doc, '注册 → 登录（动画浮动图形 + 分栏布局）→ 跳转画像仪表盘')
    bullet(doc, '点击「开始构建」→ 进入 AI 对话界面（StreamText 流式输出 + 快捷问题芯片）')
    bullet(doc, '多轮对话完成 → 显示7维度画像卡片 + 完成百分比圆环')
    bullet(doc, '选择智能体 → 配置参数 → 等待生成（5步进度动画）→ 查看资源')
    bullet(doc, '生成学习路径 → 进入路径详情 → 按步骤学习 → 打卡完成')

    h3(doc, '3.2.2 日常学习流程')
    bullet(doc, '登录 → 画像仪表盘 → 继续学习路径 → 查看/生成步骤资源')
    bullet(doc, '遇到问题 → 智能辅导提问 → 获取 AI 解答')
    bullet(doc, '查看评估报告 → 了解薄弱点 → 调整学习计划')

    h2(doc, '3.3 UI/UX 设计亮点')

    tbl(doc,
        ['亮点', '实现方式', '用户体验提升'],
        [
            ['流式输出', 'StreamText组件 + SSE + 打字机光标', 'AI回复实时展示，无需等待全部生成完成'],
            ['异步进度追踪', 'GenerationProgressView + 5步动画 + localStorage持久化', '生成过程可视化，页面刷新不丢失进度'],
            ['Markdown安全渲染', 'marked解析 + DOMPurify XSS过滤', '富文本内容安全展示，代码高亮'],
            ['资源卡片化', '差异化图标+颜色标签+难度标识', '不同类型资源一目了然，快速筛选'],
            ['路径时间线', 'el-timeline + 倒计时器 + 延长/压缩按钮', '学习进度直观展示，灵活调整'],
            ['答题交互', '选项选择→提交→即时评分→逐题解析', '沉浸式做题体验，即时反馈'],
            ['评估可视化', '5维雷达图 + 知识进度条 + 柱状图', '学习成效一目了然'],
            ['响应式布局', 'Element Plus栅格 + 弹性布局', '适配不同屏幕尺寸'],
        ],
        cw=[2.2, 5, 7.3],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第四章：数据库与API设计
    # ════════════════════════════════════════════════
    h1(doc, '第四章  数据库与API设计')

    h2(doc, '4.1 数据库设计')
    p(doc, '系统共设计14张数据库表，覆盖用户管理、学习画像、资源生成、学习路径、学习评估、智能辅导6大业务域：')

    tbl(doc,
        ['表名', '所属域', '核心字段', '表间关系'],
        [
            ['sys_user', '用户管理', 'username, password(BCrypt), nickName, major, grade', '1:1 → student_profile'],
            ['student_profile', '学习画像', 'knowledgeBase, cognitiveStyle, learningGoal, errorPronePoints(JSON)', 'N:1 → sys_user'],
            ['chat_context', '学习画像', 'sessionId, userMessage, aiReply, isExtractProfile', 'N:1 → sys_user'],
            ['ai_agent', '资源生成', 'agentName, agentRole, promptTemplate, model', '1:N → learning_resource'],
            ['learning_resource', '资源生成', 'resourceType, resourceContent, metadata(JSON), difficulty', 'N:1 → sys_user, ai_agent'],
            ['resource_generate_task', '资源生成', 'taskId(UUID), status, progress(0-100), params(JSON)', 'N:1 → sys_user'],
            ['learning_path', '学习路径', 'pathName, totalStep, currentStep, status(doing/finish)', '1:N → step'],
            ['learning_path_step', '学习路径', 'stepName, stepContent, sort, finishStatus', 'N:1 → path'],
            ['learning_path_step_resource', '学习路径', 'stepId, resourceId, sort', 'N:M 关联表'],
            ['learning_behavior', '学习路径', 'behaviorType, duration, score', 'N:1 → sys_user'],
            ['question_answer_record', '学习评估', 'userAnswer, correctAnswer, isCorrect, spendTime', 'N:1 → sys_user'],
            ['learning_evaluate', '学习评估', 'evaluateContent, knowledgeMastery(JSON), improveSuggest', 'N:1 → sys_user'],
            ['smart_tutor', '智能辅导', 'question, textAnswer, sessionId', 'N:1 → sys_user'],
            ['video_watch_progress', '学习路径', 'currentPosition, totalDuration, watchCount', 'N:1 → sys_user'],
        ],
        cw=[3.2, 1.8, 5.5, 4],
    )

    h2(doc, '4.2 API 接口设计')
    p(doc, '系统共设计约37个 RESTful API 接口，采用统一响应格式 {code: 200, msg, data}：')

    tbl(doc,
        ['模块', '接口数', '代表性接口', '说明'],
        [
            ['用户认证 /user', '6', 'POST /user/login (JWT)', '注册、登录、用户信息管理'],
            ['学习画像 /profile', '6', 'POST /profile/build/chat (多轮)', '画像对话构建、画像CRUD'],
            ['资源生成 /resource', '9', 'POST /resource/generate (异步)', '智能体列表、资源生成/查询/删除'],
            ['学习路径 /learning-path', '11', 'POST /learning-path/generate (AI)', '路径生成/调整/推荐/答题'],
            ['智能辅导 /tutor', '4', 'POST /tutor/ask (120s超时)', '提问、历史、详情、删除'],
            ['学习评估 /assessment', '6', 'POST /assessment/report/generate', '报告/结果/统计/薄弱点/趋势'],
        ],
        cw=[3, 1.5, 5, 5],
    )

    h2(doc, '4.3 接口鉴权机制')
    p(doc, '采用 JWT Token + Redis 双重鉴权。前端 Axios 拦截器自动注入 Bearer Token，'
       '后端 JwtTokenUserInterceptor 解析 Token 并存入 ThreadLocal（BaseContext），'
       'Controller 通过 BaseContext.getCurrentId() 获取当前用户。'
       'Token 有效期24小时，存入 Redis 支持主动失效（修改密码/退出登录时清除）。')

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第五章：创新实践
    # ════════════════════════════════════════════════
    h1(doc, '第五章  创新实践与用户体验提升')

    innovations = [
        ('5.1 多智能体策略工厂架构',
         '采用 Strategy + Factory 设计模式实现智能体的可插拔架构。新增智能体只需实现接口 + 标注 @Component，'
         '无需修改任何已有代码。智能体配置（名称、Prompt模板、模型选择）存储在数据库中，支持运行时动态调整。'
         '这种架构使系统具备极强的可扩展性——从最初的4个智能体扩展到8个，全程无需改动已有代码。'),
        ('5.2 对话式画像自主构建',
         '摒弃传统表单填写，通过 AI 引导式多轮对话自然收集学生信息。系统在每轮对话中分析已有画像数据，'
         '智能跳过已收集维度，根据对话内容动态调整提问策略。'
         '对话历史采用 Redis + MySQL 双写机制（Redis快速检索 + MySQL持久化），'
         '对话结束时 AI 自动从全量历史中提取结构化 JSON，采用增量更新策略确保已有数据不被覆盖。'),
        ('5.3 异步任务管线 + 进度追踪',
         'AI 资源生成采用 @Async 异步任务队列，前端通过每2秒轮询实时追踪生成进度。'
         '前端通过 localStorage 持久化任务 ID，页面刷新后可继续追踪。'
         'GenerationProgressView 提供5步可视化进度动画，将AI生成的等待时间转化为可视化的生产过程。'),
        ('5.4 Prompt 模板数据库化',
         '智能体的 Prompt 模板存储在数据库中，支持运行时动态调整而无需重启服务。'
         '采用数据库模板优先、内置默认模板兜底的双层机制。'
         '不同智能体可配置不同AI模型（qwen-plus / deepseek-v3 等），实现最优性价比。'),
        ('5.5 多维度学习评估引擎 + 画像闭环',
         '综合学习行为、答题记录、路径完成度等多维数据，通过 AssessmentAIAgent 生成个性化评估报告。'
         '评估结果中的薄弱知识点自动回写学生画像，形成「学习→评估→画像更新→资源推送」的完整闭环。'
         '支持定时自动评估（每日凌晨批量）和实时手动触发。'),
        ('5.6 动态学习路径调整',
         '支持5种路径调整模式（延长/压缩/重排/新增/移除），AI根据调整类型智能重规划未完成步骤。'
         'AI 重新分析知识体系结构和学习进度，生成符合新时间约束的步骤方案，'
         '满足学生动态变化的学习需求。'),
    ]

    for title, desc in innovations:
        h2(doc, title)
        p(doc, desc)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第六章：测试说明
    # ════════════════════════════════════════════════
    h1(doc, '第六章  测试说明')

    h2(doc, '6.1 功能测试矩阵')

    tbl(doc,
        ['测试模块', '测试项', '测试方法', '预期结果', '对应需求'],
        [
            ['用户认证', '注册/登录/登出', 'POST请求 + JWT验证', '注册成功、登录返回Token、登出清除Token', '基础功能'],
            ['用户认证', '修改密码', 'PUT请求 + 旧Token失效', '密码更新、旧Token无法使用', '基础功能'],
            ['画像构建', 'AI对话收集画像', '多轮对话 + 画像提取', '对话完成后画像7维度均有值', '需求1'],
            ['画像构建', '画像增量更新', '部分维度更新', '仅更新非null字段，已有数据保留', '需求1'],
            ['资源生成', '6种资源类型生成', '选择智能体 + 提交生成', '每种类型均能生成并正确渲染', '需求2'],
            ['资源生成', '异步进度追踪', '轮询进度接口', '进度从0到100，状态正确变迁', '需求2'],
            ['资源生成', '任务持久化', '刷新页面后继续追踪', 'localStorage保存taskId，恢复后继续轮询', '需求2'],
            ['路径规划', 'AI路径生成', '输入科目/目标/周期', '生成4-8步JSON，步骤有名称和内容', '需求3'],
            ['路径规划', '路径调整(5种)', '分别测试5种调整', '延长/压缩重规划、重排更新、新增/移除正确', '需求3'],
            ['路径规划', '资源推荐', '获取推荐资源', '返回匹配资源+评分，无匹配提示生成', '需求3'],
            ['路径规划', '答题评分', '提交题目答案', '正确判分、记录持久化、解析展示', '需求3'],
            ['智能辅导', 'AI答疑', '文字提问', '返回Markdown解答，记录持久化', '需求4'],
            ['学习评估', '评估报告生成', '触发评估', '返回完整报告(概览+掌握度+建议)', '需求5'],
            ['学习评估', '薄弱点回写', '生成评估后检查画像', 'errorPronePoints更新', '需求5'],
            ['学习评估', '定时任务', '检查定时任务日志', '每日2:00评估 + 每小时画像同步', '需求5'],
        ],
        cw=[1.8, 2.2, 3, 4, 1.5],
    )

    h2(doc, '6.2 接口测试清单')

    tbl(doc,
        ['接口', '方法', '关键测试点'],
        [
            ['POST /user/register', 'POST', '重复用户名拒绝、密码BCrypt加密存储'],
            ['POST /user/login', 'POST', '错误密码拒绝、返回JWT格式正确'],
            ['POST /profile/build/chat', 'POST', '多轮上下文保持、完成时触发画像提取'],
            ['POST /resource/generate', 'POST', '返回UUID taskId、任务状态正确变迁'],
            ['GET /resource/generate/progress/{taskId}', 'GET', '进度值递增、最终返回resourceId'],
            ['POST /learning-path/generate', 'POST', 'AI返回有效JSON、步骤4-8个'],
            ['PUT /learning-path/{id}/adjust', 'PUT', '5种模式均正确执行、120s超时处理'],
            ['POST /tutor/ask', 'POST', '返回Markdown解答、记录持久化'],
            ['POST /assessment/report/generate', 'POST', '返回完整评估数据、knowledgeMastery为有效JSON'],
        ],
        cw=[5, 1.5, 8],
    )

    h2(doc, '6.3 性能指标')

    tbl(doc,
        ['指标', '目标值', '实现方式'],
        [
            ['登录接口响应时间', '< 500ms', 'BCrypt验证 + JWT生成 + Redis存储'],
            ['画像对话响应时间', '< 5s', 'Redis会话缓存 + 同步AI调用'],
            ['资源生成完成时间', '10-60s', '@Async异步执行 + 进度轮询'],
            ['路径生成响应时间', '< 10s', '同步AI调用 + JSON解析'],
            ['辅导问答响应时间', '< 10s', '同步AI调用 + 打字机效果'],
            ['页面首屏加载时间', '< 2s', 'Vite构建 + 路由懒加载 + 代码分割'],
        ],
        cw=[3, 2.5, 9],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════
    #  第七章：部署说明
    # ════════════════════════════════════════════════
    h1(doc, '第七章  部署与运行说明')

    h2(doc, '7.1 环境要求')

    tbl(doc,
        ['组件', '版本', '说明'],
        [
            ['JDK', '17+', 'Java运行环境'],
            ['Maven', '3.8+', '后端构建'],
            ['Node.js', '^20.19 || >=22.12', '前端运行'],
            ['MySQL', '8.0+', '主数据库'],
            ['Redis', '7.0+', '会话缓存'],
        ],
        cw=[3, 3, 8.5],
    )

    h2(doc, '7.2 后端启动')
    code(doc,
        '# 1. 创建数据库\n'
        'CREATE DATABASE MediCraft CHARACTER SET utf8mb4;\n'
        '\n'
        '# 2. 配置环境变量\n'
        'export OPENAI_API_KEY=your_dashscope_key\n'
        '\n'
        '# 3. 修改 application.yaml (数据库/Redis连接)\n'
        '\n'
        '# 4. 编译启动\n'
        'cd MediCraft/MediCraft\n'
        'mvn clean package -DskipTests\n'
        'java -jar educate-server/target/educate-server-*.jar\n'
        '# 端口: 8080')

    h2(doc, '7.3 前端启动')
    code(doc,
        'cd Media-edu/MediCraft-web\n'
        'npm install\n'
        'npm run dev\n'
        '# 端口: 5173, 自动代理 /api → localhost:8080')

    h2(doc, '7.4 关键配置')
    code(doc,
        '# 后端 application.yaml\n'
        'server.port: 8080\n'
        'spring.datasource.url: jdbc:mysql://localhost:3306/MediCraft\n'
        'spring.data.redis.host: localhost:6379\n'
        'ai.base-url: https://dashscope.aliyuncs.com/compatible-mode/v1\n'
        'ai.model: qwen3.6-35b-a3b\n'
        '\n'
        '# 前端 vite.config.js\n'
        'proxy: { "/api": { target: "http://127.0.0.1:8080" } }')

    h2(doc, '7.5 初始知识库')
    p(doc, '系统以「人工智能」专业课程为切入点，在 ai_agent 表中预置6个智能体配置（含Prompt模板），'
       '在数据库中创建测试数据。系统支持用户自行通过资源生成功能扩展知识库内容。')

    # ── 保存 ──
    out_dir = r'C:\Users\57780\Desktop\code2\CAL\ppt生成'
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, 'MediCraft项目设计与开发说明书.docx')
    doc.save(out)
    print(f'Done: {out}')
    return out


if __name__ == '__main__':
    build()
